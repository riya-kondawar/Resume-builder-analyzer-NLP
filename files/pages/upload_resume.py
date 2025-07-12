import streamlit as st

# Configure page
# st.set_page_config(page_title="Resume Analyzer Pro", layout="wide")

import PyPDF2
import re
import nltk
from nltk.corpus import stopwords
from io import BytesIO
from collections import Counter
import spacy
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from wordcloud import WordCloud
from sklearn.feature_extraction.text import TfidfVectorizer
import plotly.express as px
from dateutil.relativedelta import relativedelta


# Load English language model for spaCy
try:
    nlp = spacy.load("en_core_web_sm")
except:
    st.error("Please install spaCy and its English model: `python -m spacy download en_core_web_sm`")
    nlp = None

# Download NLTK data if not present
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
    nltk.data.find('averaged_perceptron_tagger')
except:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('averaged_perceptron_tagger')

def upload_resume():
    """Enhanced Resume Analysis Page with Advanced Features"""
    st.title("📄 Advanced Resume Analyzer Pro")
    st.markdown("""
    Upload your resume for comprehensive analysis including:
    - Skill extraction and visualization
    - Experience level detection with timeline
    - Education details with institution recognition
    - Keyword optimization and ATS scoring
    - Job matching potential
    - Resume quality metrics
    """)
    
    uploaded_file = st.file_uploader(
        "Upload Resume (PDF, DOCX, TXT)", 
        type=["pdf", "docx", "txt"],
        help="Maximum file size: 200MB"
    )
    
    if uploaded_file is not None:
        try:
            file_type = uploaded_file.type
            text = ""
            
            if file_type == "application/pdf":
                text = extract_text_from_pdf(uploaded_file)
            elif file_type == "text/plain":
                text = uploaded_file.read().decode("utf-8")
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = extract_text_from_docx(uploaded_file)
            else:
                st.error("Unsupported file format")
                return
                
            # Display file info and extracted text
            with st.expander("📄 File Information & Extracted Text", expanded=True):
                col1, col2 = st.columns([1, 3])
                col1.metric("File Size", f"{len(uploaded_file.getvalue()) / 1024:.1f} KB")
                col1.metric("File Type", uploaded_file.type.split('/')[-1].upper())
                
                col2.subheader("Extracted Text Preview")
                col2.text_area("", text[:2000] + "..." if len(text) > 2000 else text, 
                             height=200, label_visibility="collapsed")
                
                if st.checkbox("Show Full Extracted Text"):
                    st.text(text)
            
            if text:
                analyze_resume(text)
                
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")

def extract_text_from_pdf(pdf_file):
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_file.read()))
        for page in pdf_reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
    except Exception as e:
        raise Exception(f"PDF extraction error: {str(e)}")
    return text.strip()

def extract_text_from_docx(docx_file):
    try:
        from docx import Document
        doc = Document(BytesIO(docx_file.read()))
        return "\n".join([para.text for para in doc.paragraphs])
    except:
        st.warning("DOCX parsing requires python-docx library. Install with: pip install python-docx")
        return ""

def analyze_resume(text):
    with st.spinner("Performing deep analysis... This may take a moment for comprehensive reports."):
        # Initialize tabs for different analysis sections
        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "📊 Overview", 
            "🛠 Skills Analysis", 
            "📅 Experience Analysis", 
            "🎓 Education Analysis", 
            "🔍 Advanced Metrics"
        ])
        
        with tab1:
            display_overview_analysis(text)
        
        with tab2:
            display_skills_analysis(text)
        
        with tab3:
            display_experience_analysis(text)
        
        with tab4:
            display_education_analysis(text)
        
        with tab5:
            display_advanced_metrics(text)

def display_overview_analysis(text):
    """Display comprehensive overview analysis"""
    st.subheader("📊 Resume Overview Metrics")
    
    # Basic Statistics
    col1, col2, col3, col4 = st.columns(4)
    word_count = len(text.split())
    char_count = len(text)
    sentence_count = len(nltk.sent_tokenize(text))
    unique_word_count = len(set(word.lower() for word in text.split()))
    
    col1.metric("Word Count", word_count)
    col2.metric("Character Count", char_count)
    col3.metric("Sentence Count", sentence_count)
    col4.metric("Unique Words", unique_word_count)
    
    # Calculate readability scores
    flesch_score = calculate_flesch_reading_ease(text)
    grade_level = calculate_flesch_kincaid_grade(text)
    
    col1, col2 = st.columns(2)
    col1.metric("Readability Score", f"{flesch_score:.1f}/100", 
               help="Higher scores (60-100) indicate easier reading")
    col2.metric("Grade Level", f"Grade {grade_level:.1f}", 
               help="Lower grade levels are easier to read")
    
    # Word Cloud Visualization
    st.subheader("🔠 Keyword Cloud")
    generate_wordcloud(text)
    
    # Sentiment Analysis
    sentiment_score, sentiment_label = analyze_sentiment(text)
    col1, col2 = st.columns(2)
    col1.metric("Overall Sentiment", sentiment_label, 
               help="Positive sentiment indicates confident, achievement-oriented language")
    col2.metric("Sentiment Score", f"{sentiment_score:.2f}/5.0")
    
    # Keyword Analysis
    st.subheader("🔑 Top Keywords")
    keywords = extract_keywords(text)
    display_keyword_bars(keywords[:20])

def display_skills_analysis(text):
    """Display skills analysis with visualizations"""
    st.subheader("🛠 Technical Skills Analysis")
    
    skills = extract_skills(text)
    soft_skills = extract_soft_skills(text)
    
    if not skills and not soft_skills:
        st.warning("No skills detected in the resume. Consider adding a skills section.")
        return
    
    # Skills Visualization
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Technical Skills**")
        if skills:
            skills_df = pd.DataFrame(skills, columns=["Skill"])
            skills_df["Category"] = skills_df["Skill"].apply(categorize_skill)
            
            fig = px.treemap(skills_df, path=["Category", "Skill"], 
                            title="Technical Skills Breakdown")
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.warning("No technical skills detected")
    
    with col2:
        st.markdown("**Soft Skills**")
        if soft_skills:
            soft_skills_df = pd.DataFrame(soft_skills, columns=["Skill"])
            soft_skills_df["Count"] = 1  # Placeholder for frequency
            
            fig = px.bar(soft_skills_df, x="Skill", y="Count", 
                         title="Soft Skills Frequency")
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.warning("No soft skills detected")
    
    # Skills Gap Analysis
    st.subheader("🔍 Skills Gap Analysis")
    target_job = st.text_input("Enter a target job title for skills comparison:", 
                              "Software Engineer")
    
    if target_job and skills:
        recommended_skills = get_recommended_skills(target_job)
        missing_skills = set(recommended_skills) - set(skills)
        
        if missing_skills:
            st.warning(f"Your resume is missing these skills commonly required for {target_job}:")
            cols = st.columns(4)
            for i, skill in enumerate(missing_skills):
                cols[i%4].error(skill)
        else:
            st.success("Your resume contains all key skills typically required for this role!")

def display_experience_analysis(text):
    """Display experience analysis with timeline"""
    st.subheader("📅 Professional Experience Analysis")
    
    experience = detect_experience(text)
    
    # Experience Summary
    col1, col2, col3 = st.columns(3)
    col1.metric("Total Experience", f"{experience['total_years']} years")
    col2.metric("Position Count", len(experience['positions']))
    col3.metric("Current Position", 
               experience['positions'][0]['title'] if experience['positions'] else "N/A")
    
    # Experience Timeline
    if experience['positions']:
        st.subheader("⏳ Work History Timeline")
        
        # Prepare timeline data
        timeline_data = []
        for pos in experience['positions']:
            start_year = int(pos['start'])
            end_year = int(pos['end']) if pos['end'].isdigit() else datetime.now().year
            timeline_data.append({
                "Position": pos['title'],
                "Company": pos['company'],
                "Start": start_year,
                "End": end_year,
                "Duration": end_year - start_year
            })
        
        timeline_df = pd.DataFrame(timeline_data)
        
        # Create timeline visualization
        fig = px.timeline(
            timeline_df, 
            x_start="Start", 
            x_end="End", 
            y="Position",
            color="Company",
            title="Professional Experience Timeline",
            hover_name="Company",
            hover_data={"Duration": True, "Start": False, "End": False}
        )
        fig.update_yaxes(autorange="reversed")
        st.plotly_chart(fig, use_container_width=True)
        
        # Display positions table
        st.dataframe(timeline_df.sort_values("Start", ascending=False).reset_index(drop=True))
    else:
        st.warning("No work experience detected in the resume")

def display_education_analysis(text):
    """Display education analysis"""
    st.subheader("🎓 Education Background Analysis")
    
    education = detect_education(text)
    
    if education:
        # Education Timeline
        edu_df = pd.DataFrame(education)
        edu_df['Year'] = pd.to_numeric(edu_df['year'], errors='coerce')
        edu_df = edu_df.dropna(subset=['Year'])
        
        if not edu_df.empty:
            fig = px.scatter(
                edu_df,
                x="Year",
                y="degree",
                color="institution",
                size=[10]*len(edu_df),
                title="Education Timeline",
                labels={"degree": "Degree", "institution": "Institution"}
            )
            st.plotly_chart(fig, use_container_width=True)
        
        # Education Details
        st.dataframe(edu_df.sort_values("Year", ascending=False).reset_index(drop=True))
        
        # Education Quality Metrics
        highest_degree = edu_df.loc[edu_df['Year'].idxmax()]['degree'] if not edu_df.empty else "N/A"
        institution_rank = estimate_institution_rank(edu_df['institution'].tolist()[0]) if not edu_df.empty else "N/A"
        
        col1, col2 = st.columns(2)
        col1.metric("Highest Degree", highest_degree)
        col2.metric("Institution Ranking Estimate", institution_rank)
    else:
        st.warning("No education information detected in the resume")

def display_advanced_metrics(text):
    """Display advanced resume metrics"""
    st.subheader("🧮 Advanced Resume Metrics")
    
    # ATS Optimization Score
    ats_score = calculate_ats_score(text)
    st.metric("ATS Optimization Score", f"{ats_score:.1f}/100",
             help="Score based on how well your resume is optimized for Applicant Tracking Systems")
    
    # Keyword Density Analysis
    st.subheader("📈 Keyword Density Analysis")
    keywords = extract_keywords(text)
    keyword_density = {word: count/len(keywords) for word, count in Counter(keywords).items()}
    top_keywords = dict(sorted(keyword_density.items(), key=lambda item: item[1], reverse=True)[:10])
    
    fig, ax = plt.subplots()
    ax.barh(list(top_keywords.keys()), list(top_keywords.values()))
    ax.set_xlabel("Density")
    ax.set_title("Top Keywords by Density")
    st.pyplot(fig)
    
    # Action Verb Analysis
    st.subheader("🏃 Action Verbs Analysis")
    action_verbs = extract_action_verbs(text)
    if action_verbs:
        cols = st.columns(4)
        for i, verb in enumerate(action_verbs[:20]):
            cols[i%4].success(verb)
    else:
        st.warning("Few action verbs detected. Consider adding more achievement-oriented language.")
    
    # Resume Length Analysis
    st.subheader("📏 Resume Length Analysis")
    word_count = len(text.split())
    ideal_range = (400, 800)  # Words
    
    if word_count < ideal_range[0]:
        st.error(f"Resume may be too short ({word_count} words). Consider adding more details.")
    elif word_count > ideal_range[1]:
        st.warning(f"Resume may be too long ({word_count} words). Consider being more concise.")
    else:
        st.success(f"Resume length is ideal ({word_count} words).")

# ========== Analysis Helper Functions ==========

def extract_skills(text):
    """Enhanced skill extraction with more categories and patterns"""
    skill_patterns = {
        'Programming': r'\b(python|java|c\+\+|c#|javascript|typescript|go|rust|ruby|php|swift|kotlin|scala|r\b|perl)\b',
        'Web': r'\b(html|css|sass|less|react|angular|vue|django|flask|node\.?js|express|spring|laravel|rails|asp\.net)\b',
        'Mobile': r'\b(android|ios|flutter|react native|xamarin|ionic)\b',
        'Data': r'\b(sql|nosql|mysql|postgresql|mongodb|redis|cassandra|hadoop|spark|pandas|numpy|tensorflow|pytorch|scikit-learn|keras)\b',
        'Cloud': r'\b(aws|azure|gcp|docker|kubernetes|terraform|ansible|jenkins|ci/cd|devops|serverless)\b',
        'Data Science': r'\b(machine learning|deep learning|nlp|computer vision|data mining|data visualization|tableau|power bi)\b',
        'Other Tech': r'\b(git|linux|bash|rest api|graphql|microservices|blockchain|cybersecurity|ethical hacking)\b'
    }
    
    skills = set()
    text_lower = text.lower()
    
    for category, pattern in skill_patterns.items():
        matches = re.findall(pattern, text_lower)
        skills.update(matches)
    
    return sorted(skills)

def extract_soft_skills(text):
    """Extract soft skills from resume text"""
    soft_skills = [
        'communication', 'leadership', 'teamwork', 'problem solving', 
        'creativity', 'time management', 'adaptability', 'critical thinking',
        'collaboration', 'emotional intelligence', 'negotiation', 'conflict resolution',
        'decision making', 'strategic thinking', 'mentoring', 'public speaking',
        'presentation', 'interpersonal', 'analytical', 'attention to detail'
    ]
    
    found_skills = []
    text_lower = text.lower()
    
    for skill in soft_skills:
        if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
            found_skills.append(skill.title())
    
    return found_skills

def categorize_skill(skill):
    """Categorize a skill into broader categories"""
    skill = skill.lower()
    categories = {
        'Programming': ['python', 'java', 'c++', 'c#', 'javascript', 'typescript', 'go', 'rust', 'ruby', 'php', 'swift', 'kotlin', 'scala', 'r', 'perl'],
        'Web': ['html', 'css', 'sass', 'less', 'react', 'angular', 'vue', 'django', 'flask', 'node', 'express', 'spring', 'laravel', 'rails', 'asp.net'],
        'Mobile': ['android', 'ios', 'flutter', 'react native', 'xamarin', 'ionic'],
        'Data': ['sql', 'nosql', 'mysql', 'postgresql', 'mongodb', 'redis', 'cassandra', 'hadoop', 'spark', 'pandas', 'numpy', 'tensorflow', 'pytorch', 'scikit-learn', 'keras'],
        'Cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'ansible', 'jenkins', 'ci/cd', 'devops', 'serverless'],
        'Data Science': ['machine learning', 'deep learning', 'nlp', 'computer vision', 'data mining', 'data visualization', 'tableau', 'power bi'],
        'Other': ['git', 'linux', 'bash', 'rest api', 'graphql', 'microservices', 'blockchain', 'cybersecurity', 'ethical hacking']
    }
    
    for category, skills in categories.items():
        if any(s in skill for s in skills):
            return category
    return 'Other'

def detect_experience(text):
    """Enhanced experience detection with better pattern matching"""
    experience = {
        'total_years': 0,
        'positions': []
    }
    
    # Improved experience pattern matching
    exp_patterns = [
        r'(\d+)\+?\s*(years?|yrs?)\s*(of)?\s*(experience|exp)',
        r'(\d+)\s*\+\s*(years?|yrs?)\s*(of)?\s*(experience|exp)',
        r'(\d+)\s*-\s*(\d+)\s*(years?|yrs?)\s*(of)?\s*(experience|exp)'
    ]
    
    for pattern in exp_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            if match.group(1):
                years = int(match.group(1))
                if years > experience['total_years']:
                    experience['total_years'] = years
    
    # Enhanced position detection
    position_patterns = [
        r'(?P<title>\b[A-Z][a-zA-Z\s]+\b)\s+(?:at|@)\s+(?P<company>\b[A-Z][a-zA-Z\s&]+\b)\s*\(?(?P<start>\d{4})\s*(?:-|to)\s*(?P<end>\d{4}|present|current)\)?',
        r'(?P<title>\b[A-Z][a-zA-Z\s]+\b)\s*,\s*(?P<company>\b[A-Z][a-zA-Z\s&]+\b)\s*,\s*(?P<start>\w+\s+\d{4})\s*(?:-|to)\s*(?P<end>\w+\s+\d{4}|present|current)',
        r'(?P<company>\b[A-Z][a-zA-Z\s&]+\b)\s*-\s*(?P<title>\b[A-Z][a-zA-Z\s]+\b)\s*,\s*(?P<start>\w+\s+\d{4})\s*(?:-|to)\s*(?P<end>\w+\s+\d{4}|present|current)'
    ]
    
    for pattern in position_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            start_year = match.group('start')
            if len(start_year) > 4:  # If month is included
                try:
                    start_date = datetime.strptime(start_year, "%B %Y")
                    start_year = start_date.year
                except:
                    start_year = re.search(r'\d{4}', start_year).group()
            
            end_year = match.group('end')
            if not end_year.isdigit():
                end_year = datetime.now().year
            
            experience['positions'].append({
                'title': match.group('title').strip(),
                'company': match.group('company').strip(),
                'start': str(start_year),
                'end': str(end_year)
            })
    
    return experience

def detect_education(text):
    """Enhanced education detection with degree types and institutions"""
    education = []
    
    # Improved degree patterns
    degree_patterns = [
        r'\b(B\.?[A-Za-z]*\.?|Bachelor\s*of\s*[A-Za-z]+|M\.?[A-Za-z]*\.?|Master\s*of\s*[A-Za-z]+|PhD|Doctorate)\b',
        r'\b(Associate|Diploma|Certificate|High School|GED)\b'
    ]
    
    # Institution patterns
    institution_patterns = [
        r'(?:University|College|Institute|School)\s+of\s+[A-Za-z\s&]+',
        r'[A-Z][a-zA-Z\s&]+(?:University|College|Institute|School)',
        r'[A-Z]{2,}\b'  # For acronyms like MIT, UCLA
    ]
    
    # Year patterns
    year_pattern = r'(?:19|20)\d{2}'
    
    # Find education sections with more flexible matching
    edu_sections = re.finditer(
        r'(?:education|academics|qualifications)[^a-z0-9]*(.*?)(?:(?:\n\n)|(?=\n[A-Z][a-z])|\Z)', 
        text, 
        re.IGNORECASE | re.DOTALL
    )
    
    for section in edu_sections:
        section_text = section.group(1)
        
        # Extract degrees
        degrees = []
        for pattern in degree_patterns:
            degrees.extend([m.group() for m in re.finditer(pattern, section_text, re.IGNORECASE)])
        
        # Extract institutions
        institutions = []
        for pattern in institution_patterns:
            institutions.extend([m.group() for m in re.finditer(pattern, section_text)])
        
        # Extract years
        years = re.findall(year_pattern, section_text)
        
        # Pair them up
        for i, degree in enumerate(degrees):
            education.append({
                'degree': degree,
                'institution': institutions[i] if i < len(institutions) else "Unknown",
                'year': years[i] if i < len(years) else "Unknown"
            })
    
    return education

def extract_keywords(text):
    """Enhanced keyword extraction using TF-IDF and POS tagging"""
    stop_words = set(stopwords.words('english'))
    custom_stopwords = {'experience', 'work', 'project', 'skill', 'education', 'university'}
    stop_words.update(custom_stopwords)
    
    # Tokenize and filter
    words = nltk.word_tokenize(text.lower())
    words = [word for word in words if word.isalpha() and word not in stop_words and len(word) > 2]
    
    # POS tagging to focus on nouns and adjectives
    tagged_words = nltk.pos_tag(words)
    nouns_and_adjs = [word for word, pos in tagged_words if pos.startswith('NN') or pos.startswith('JJ')]
    
    # Count frequency
    word_freq = Counter(nouns_and_adjs)
    
    # Get top keywords
    return [word for word, count in word_freq.most_common(50)]

def extract_action_verbs(text):
    """Extract action verbs from resume text"""
    verbs = [
        'achieved', 'managed', 'led', 'developed', 'implemented', 'designed',
        'created', 'improved', 'increased', 'reduced', 'optimized', 'solved',
        'built', 'launched', 'initiated', 'established', 'transformed', 'mentored',
        'trained', 'coordinated', 'organized', 'presented', 'published', 'authored',
        'researched', 'analyzed', 'evaluated', 'recommended', 'negotiated', 'collaborated'
    ]
    
    found_verbs = []
    text_lower = text.lower()
    
    for verb in verbs:
        if re.search(r'\b' + re.escape(verb) + r'\b', text_lower):
            found_verbs.append(verb.title())
    
    return found_verbs

def analyze_sentiment(text):
    """Enhanced sentiment analysis with more nuanced scoring"""
    positive_words = {
        'excellent': 2, 'achieved': 1.5, 'success': 1.5, 'lead': 1, 
        'improved': 1.5, 'award': 2, 'innovative': 1, 'optimized': 1,
        'increased': 1, 'reduced': 1, 'solved': 1, 'created': 1,
        'developed': 1, 'transformed': 1.5, 'outstanding': 2
    }
    
    negative_words = {
        'poor': -2, 'lack': -1, 'failed': -1.5, 'issue': -1,
        'problem': -1, 'difficulty': -1, 'challenge': -0.5, 'limited': -1
    }
    
    # Calculate sentiment score
    positive_score = sum(
        text.lower().count(word) * weight 
        for word, weight in positive_words.items()
    )
    
    negative_score = sum(
        text.lower().count(word) * weight 
        for word, weight in negative_words.items()
    )
    
    # Normalize score to 1-5 range
    base_score = 3.0
    adjusted_score = base_score + (positive_score - negative_score) / 20
    final_score = max(1.0, min(5.0, adjusted_score))
    
    # Determine label
    if final_score >= 4.0:
        label = "Highly Positive"
    elif final_score >= 3.0:
        label = "Positive"
    elif final_score >= 2.0:
        label = "Neutral"
    else:
        label = "Negative"
    
    return final_score, label

def calculate_flesch_reading_ease(text):
    """Calculate Flesch Reading Ease score"""
    sentences = nltk.sent_tokenize(text)
    words = [nltk.word_tokenize(sentence) for sentence in sentences]
    total_sentences = len(sentences)
    total_words = sum(len(word_list) for word_list in words)
    total_syllables = sum(count_syllables(word) for word_list in words for word in word_list)
    
    if total_sentences == 0 or total_words == 0:
        return 0
    
    score = 206.835 - 1.015 * (total_words / total_sentences) - 84.6 * (total_syllables / total_words)
    return max(0, min(100, score))

def calculate_flesch_kincaid_grade(text):
    """Calculate Flesch-Kincaid Grade Level"""
    sentences = nltk.sent_tokenize(text)
    words = [nltk.word_tokenize(sentence) for sentence in sentences]
    total_sentences = len(sentences)
    total_words = sum(len(word_list) for word_list in words)
    total_syllables = sum(count_syllables(word) for word_list in words for word in word_list)
    
    if total_sentences == 0 or total_words == 0:
        return 0
    
    score = 0.39 * (total_words / total_sentences) + 11.8 * (total_syllables / total_words) - 15.59
    return max(1, min(12, score))

def count_syllables(word):
    """Approximate syllable count for a word"""
    word = word.lower()
    count = 0
    
    # Count vowel groups
    vowels = "aeiouy"
    prev_char_was_vowel = False
    
    for char in word:
        if char in vowels:
            if not prev_char_was_vowel:
                count += 1
            prev_char_was_vowel = True
        else:
            prev_char_was_vowel = False
    
    # Adjust for silent e
    if word.endswith('e'):
        count -= 1
    
    # Ensure at least one syllable
    return max(1, count)

def calculate_ats_score(text):
    """Calculate ATS optimization score based on resume best practices"""
    score = 50  # Base score
    
    # Check for important sections
    sections = [
        'experience', 'education', 'skills', 
        'summary', 'projects', 'certifications'
    ]
    
    for section in sections:
        if re.search(r'\b' + section + r'\b', text, re.IGNORECASE):
            score += 5
    
    # Check for keywords
    keywords = extract_keywords(text)
    if len(keywords) >= 10:
        score += min(20, len(keywords) / 2)
    
    # Check for action verbs
    action_verbs = extract_action_verbs(text)
    if action_verbs:
        score += min(10, len(action_verbs))
    
    # Check length
    word_count = len(text.split())
    if 400 <= word_count <= 800:
        score += 10
    elif word_count > 800:
        score -= 5
    
    return min(100, score)

def estimate_institution_rank(institution):
    """Very basic institution ranking estimation (placeholder)"""
    top_tier = ['Harvard', 'MIT', 'Stanford', 'Cambridge', 'Oxford', 'Princeton']
    second_tier = ['Yale', 'Columbia', 'Chicago', 'Berkeley', 'UCLA', 'Michigan']
    
    if any(name in institution for name in top_tier):
        return "Top Tier (Global Top 20)"
    elif any(name in institution for name in second_tier):
        return "Second Tier (Global Top 100)"
    else:
        return "Recognized Institution"

def get_recommended_skills(job_title):
    """Get recommended skills for a job title (placeholder)"""
    skill_library = {
        'software engineer': ['python', 'java', 'c++', 'javascript', 'sql', 'git', 'docker', 'aws', 'rest api', 'agile'],
        'data scientist': ['python', 'r', 'sql', 'pandas', 'numpy', 'machine learning', 'tensorflow', 'pytorch', 'data visualization', 'statistics'],
        'product manager': ['agile', 'scrum', 'product strategy', 'market research', 'user stories', 'jira', 'roadmapping', 'stakeholder management', 'presentation'],
        'devops engineer': ['aws', 'azure', 'docker', 'kubernetes', 'terraform', 'ansible', 'ci/cd', 'jenkins', 'linux', 'bash']
    }
    
    job_title_lower = job_title.lower()
    for title, skills in skill_library.items():
        if title in job_title_lower:
            return skills
    
    # Default return if job title not found
    return ['python', 'communication', 'problem solving', 'teamwork']

def generate_wordcloud(text):
    """Generate and display a word cloud"""
    wordcloud = WordCloud(
        width=800,
        height=400,
        background_color='white',
        stopwords=set(stopwords.words('english'))
    ).generate(text)
    
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.imshow(wordcloud, interpolation='bilinear')
    ax.axis('off')
    st.pyplot(fig)

def display_keyword_bars(keywords):
    """Display horizontal bar chart of keywords"""
    fig, ax = plt.subplots(figsize=(10, 6))
    y_pos = range(len(keywords))
    ax.barh(y_pos, [10]*len(keywords))  # Equal width for all
    ax.set_yticks(y_pos)
    ax.set_yticklabels(keywords)
    ax.invert_yaxis()
    ax.set_title('Top Keywords')
    st.pyplot(fig)

if __name__ == "__main__":
    upload_resume()







# import streamlit as st
# import PyPDF2
# import re
# import nltk
# from nltk.corpus import stopwords
# from io import BytesIO
# from collections import Counter
# import spacy
# from datetime import datetime
# import pandas as pd

# # Load English language model for spaCy
# try:
#     nlp = spacy.load("en_core_web_sm")
# except:
#     st.error("Please install spaCy and its English model: `python -m spacy download en_core_web_sm`")
#     nlp = None

# # Download NLTK data if not present
# try:
#     nltk.data.find('tokenizers/punkt')
#     nltk.data.find('corpora/stopwords')
# except:
#     nltk.download('punkt')
#     nltk.download('stopwords')

# def upload_resume():
#     """Enhanced Resume Analysis Page"""
#     st.title("📄 Advanced Resume Analyzer")
#     st.markdown("""
#     Upload your resume for comprehensive analysis including:
#     - Skill extraction
#     - Experience level detection
#     - Education details
#     - Keyword optimization
#     - Job matching potential
#     """)
    
#     uploaded_file = st.file_uploader(
#         "Upload Resume (PDF, DOCX, TXT)", 
#         type=["pdf", "docx", "txt"],
#         accept_multiple_files=False
#     )
    
#     if uploaded_file is not None:
#         try:
#             file_type = uploaded_file.type
#             text = ""
            
#             if file_type == "application/pdf":
#                 text = extract_text_from_pdf(uploaded_file)
#             elif file_type == "text/plain":
#                 text = uploaded_file.read().decode("utf-8")
#             elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 text = extract_text_from_docx(uploaded_file)
#             else:
#                 st.error("Unsupported file format")
#                 return
                
#             st.subheader("Extracted Text Preview")
#             st.text_area("", text[:2000] + "..." if len(text) > 2000 else text, height=200)
            
#             if text:
#                 with st.expander("Show Full Extracted Text"):
#                     st.text(text)
                
#                 analyze_resume(text)
                
#         except Exception as e:
#             st.error(f"Error processing file: {str(e)}")

# def extract_text_from_pdf(pdf_file):
#     text = ""
#     try:
#         pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_file.read()))
#         for page in pdf_reader.pages:
#             page_text = page.extract_text() or ""
#             text += page_text + "\n"
#     except Exception as e:
#         raise Exception(f"PDF extraction error: {str(e)}")
#     return text.strip()

# def extract_text_from_docx(docx_file):
#     try:
#         from docx import Document
#         doc = Document(BytesIO(docx_file.read()))
#         return "\n".join([para.text for para in doc.paragraphs])
#     except:
#         return "DOCX parsing requires python-docx library"

# def analyze_resume(text):
#     with st.spinner("Performing deep analysis..."):
#         # Basic Text Analysis
#         st.subheader("📊 Basic Statistics")
#         col1, col2, col3 = st.columns(3)
#         word_count = len(text.split())
#         char_count = len(text)
#         sentence_count = len(nltk.sent_tokenize(text))
        
#         col1.metric("Word Count", word_count)
#         col2.metric("Character Count", char_count)
#         col3.metric("Sentence Count", sentence_count)
        
#         # Advanced NLP Analysis
#         if nlp:
#             doc = nlp(text)
            
#             # Named Entity Recognition
#             st.subheader("🔍 Key Information Extraction")
            
#             # Skills Extraction
#             skills = extract_skills(text)
#             st.markdown("**🛠 Technical Skills Detected**")
#             skills_cols = st.columns(4)
#             for i, skill in enumerate(skills[:20]):
#                 skills_cols[i%4].success(skill)
            
#             # Experience Detection
#             experience = detect_experience(text)
#             st.markdown("**📅 Professional Experience**")
#             st.write(f"Total Experience: {experience['total_years']} years")
#             st.dataframe(pd.DataFrame(experience['positions']))
            
#             # Education Detection
#             education = detect_education(text)
#             st.markdown("**🎓 Education Background**")
#             for edu in education:
#                 st.write(f"- {edu['degree']} from {edu['institution']} ({edu['year']})")
            
#             # Keyword Analysis
#             st.subheader("🔑 Keyword Analysis")
#             keywords = extract_keywords(text)
#             st.write("Top 20 keywords (excluding common words):")
#             st.write(", ".join(keywords[:20]))
            
#             # Sentiment Analysis
#             sentiment = analyze_sentiment(text)
#             st.metric("Overall Sentiment Score", f"{sentiment:.2f}/5.0")
            
#         else:
#             st.warning("SpaCy model not loaded - some advanced features disabled")

# def extract_skills(text):
#     """Extract technical skills using pattern matching"""
#     skill_patterns = {
#         'programming': r'\b(python|java|c\+\+|javascript|typescript|go|rust|ruby|php|swift|kotlin)\b',
#         'web': r'\b(html|css|react|angular|vue|django|flask|node\.?js|express)\b',
#         'data': r'\b(sql|nosql|mysql|postgresql|mongodb|hadoop|spark|pandas|numpy|tensorflow|pytorch)\b',
#         'devops': r'\b(docker|kubernetes|aws|azure|gcp|ci/cd|jenkins|terraform|ansible)\b'
#     }
    
#     skills = set()
#     text_lower = text.lower()
    
#     for category, pattern in skill_patterns.items():
#         matches = re.findall(pattern, text_lower)
#         skills.update(matches)
    
#     return sorted(skills)

# def detect_experience(text):
#     """Detect work experience duration and positions"""
#     experience = {
#         'total_years': 0,
#         'positions': []
#     }
    
#     # Simple pattern matching for experience
#     exp_pattern = r'(\d+)\+?\s*(years?|yrs?)\s*(of)?\s*(experience|exp)'
#     match = re.search(exp_pattern, text, re.IGNORECASE)
#     if match:
#         experience['total_years'] = int(match.group(1))
    
#     # Detect position history
#     position_pattern = r'(\b[A-Z][a-z]+\b)\s+at\s+(\b[A-Z][a-zA-Z\s]+\b)\s+\((\d{4})\s*-\s*(\d{4}|present)\)'
#     matches = re.findall(position_pattern, text)
    
#     for match in matches:
#         experience['positions'].append({
#             'title': match[0],
#             'company': match[1],
#             'start': match[2],
#             'end': match[3]
#         })
    
#     return experience

# def detect_education(text):
#     """Detect education background"""
#     education = []
    
#     # Common degree patterns
#     degree_pattern = r'\b(B\.?Sc|B\.?Tech|B\.?E|B\.?A|M\.?Sc|M\.?Tech|M\.?A|PhD)\b'
#     institution_pattern = r'((?:University|College|Institute|School)\s+of\s+[A-Za-z\s]+)'
#     year_pattern = r'(19|20)\d{2}'
    
#     # Find education sections
#     edu_sections = re.findall(r'(education[^a-z0-9].*?)(?:\n\n|\Z)', text, re.IGNORECASE | re.DOTALL)
    
#     for section in edu_sections:
#         degrees = re.findall(degree_pattern, section, re.IGNORECASE)
#         institutions = re.findall(institution_pattern, section)
#         years = re.findall(year_pattern, section)
        
#         for i, degree in enumerate(degrees):
#             education.append({
#                 'degree': degree,
#                 'institution': institutions[i] if i < len(institutions) else "Unknown",
#                 'year': years[i] if i < len(years) else "Unknown"
#             })
    
#     return education

# def extract_keywords(text):
#     """Extract important keywords using TF-IDF like approach"""
#     stop_words = set(stopwords.words('english'))
#     words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
#     filtered_words = [word for word in words if word not in stop_words and len(word) > 3]
#     word_freq = Counter(filtered_words)
#     return [word for word, count in word_freq.most_common(50)]

# def analyze_sentiment(text):
#     """Basic sentiment analysis (placeholder for actual model)"""
#     positive_words = ['excellent', 'achieved', 'success', 'lead', 'improved', 'award']
#     negative_words = ['poor', 'lack', 'failed', 'issue', 'problem']
    
#     positive_count = sum(text.lower().count(word) for word in positive_words)
#     negative_count = sum(text.lower().count(word) for word in negative_words)
    
#     return max(1.0, min(5.0, 3.0 + (positive_count - negative_count) / 10))
